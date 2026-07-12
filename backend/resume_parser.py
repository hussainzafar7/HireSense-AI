import re
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import BinaryIO

import pdfplumber
from docx import Document

CS_DOMAIN_SKILLS = [
    "variables", "loops", "functions", "recursion", "OOP", "data types", "exception handling", "file I/O", "memory management",
    "arrays", "linked lists", "stacks", "queues", "trees", "graphs", "hash tables", "heaps", "tries", "segment trees",
    "sorting", "searching", "dynamic programming", "greedy algorithms", "divide and conquer", "graph algorithms", "BFS", "DFS", "Dijkstra", "complexity analysis", "Big O notation",
    "SQL", "normalization", "indexing", "transactions", "ACID", "NoSQL", "MongoDB", "Redis", "ER diagrams", "query optimization",
    "processes", "threads", "scheduling", "virtual memory", "file systems", "deadlock", "synchronization", "semaphores", "mutexes",
    "OSI model", "TCP/IP", "HTTP", "HTTPS", "DNS", "routing", "IP addressing", "subnetting", "firewalls", "sockets", "REST API",
    "SDLC", "Agile", "Scrum", "design patterns", "SOLID principles", "testing", "CI/CD", "Git", "code review", "refactoring",
    "HTML", "CSS", "JavaScript", "React", "Node.js", "REST", "GraphQL", "authentication", "JWT", "cookies", "sessions", "WebSockets", "responsive design", "DOM manipulation",
    "supervised learning", "unsupervised learning", "neural networks", "CNN", "RNN", "LSTM", "transformers", "overfitting", "regularization", "gradient descent", "backpropagation", "feature engineering", "model evaluation", "cross-validation", "bias-variance tradeoff",
    "PyTorch", "TensorFlow", "Keras", "attention mechanism", "BERT", "GPT", "transfer learning", "embeddings", "loss functions", "optimizers", "batch normalization", "dropout",
    "AWS", "GCP", "Azure", "Docker", "Kubernetes", "microservices", "serverless", "load balancing", "auto-scaling", "monitoring", "infrastructure as code", "Terraform",
    "encryption", "hashing", "SSL", "TLS", "OAuth", "XSS", "SQL injection", "CSRF", "penetration testing", "VPN", "zero trust",
    "Photoshop", "color theory", "vector", "raster", "UI/UX", "Figma", "wireframing", "prototyping", "typography", "design systems",
    "Android", "iOS", "React Native", "Flutter", "Dart", "app lifecycle", "push notifications", "SQLite", "API calls",
    "pandas", "numpy", "matplotlib", "seaborn", "statistical analysis", "hypothesis testing", "A/B testing", "data cleaning", "EDA", "feature selection", "dimensionality reduction", "PCA",
]
EXTRA_SKILLS = [
    "Python", "Java", "C++", "C", "TypeScript", "Go", "Rust", "PHP", "Ruby", "Swift", "Kotlin", "R", "MATLAB", "Scala",
    "Django", "FastAPI", "Spring Boot", "Laravel", "Express.js", "Next.js", "Vue.js", "Angular",
    "GitHub", "Jenkins", "Jira", "Postman", "Linux", "Bash", "MySQL", "PostgreSQL", "Cassandra", "Elasticsearch",
]
MASTER_SKILLS = sorted(set(CS_DOMAIN_SKILLS + EXTRA_SKILLS), key=lambda x: x.lower())

_nlp = None


def load_spacy():
    global _nlp
    if _nlp is not None:
        return _nlp
    try:
        import spacy
        try:
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=False, timeout=180)
            _nlp = spacy.load("en_core_web_sm")
    except Exception:
        _nlp = False
    return _nlp


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or " ").strip()


def extract_pdf(path: Path) -> str:
    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return "\n".join(chunks)


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(chunks)


def text_from_upload(file_obj: BinaryIO, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(file_obj.read())
        tmp_path = Path(tmp.name)
    try:
        if suffix == ".pdf":
            text = extract_pdf(tmp_path)
        elif suffix == ".docx":
            text = extract_docx(tmp_path)
        else:
            raise ValueError("Only PDF and DOCX resumes are supported")
    finally:
        try:
            tmp_path.unlink()
        except OSError:
            pass
    if not clean_text(text):
        raise ValueError("No extractable text found. Scanned image resumes are not supported.")
    return text


def extract_name(lines: list[str], text: str) -> str:
    nlp = load_spacy()
    first_lines = "\n".join(lines[:5])
    if nlp:
        doc = nlp(first_lines)
        for ent in doc.ents:
            if ent.label_ == "PERSON" and 1 <= len(ent.text.split()) <= 4:
                return ent.text.strip()
    for line in lines[:5]:
        line = line.strip()
        if re.match(r"^[A-Z][A-Za-z.'-]+\s+[A-Z][A-Za-z.'-]+(?:\s+[A-Z][A-Za-z.'-]+)?$", line):
            return line
    return lines[0].strip() if lines else "Unknown Candidate"


def extract_sections(text: str, heading: str) -> list[str]:
    pattern = rf"(?is)(?:{heading})\s*[:\-]?\s*(.*?)(?:\n\s*(?:education|experience|projects|certifications|skills|summary)\s*[:\-]?|$)"
    match = re.search(pattern, text)
    if not match:
        return []
    content = match.group(1) or ""
    return [clean_text(part) for part in re.split(r"\n|•|- ", content) if clean_text(part)]


def parse_resume_text(text: str) -> dict:
    raw_text = text
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    compact = clean_text(text)
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", compact)
    phone_match = re.search(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}", compact)
    lower = compact.lower()
    skills = []
    for skill in MASTER_SKILLS:
        s = skill.lower()
        pattern = r"(?<![a-z0-9+#])" + re.escape(s) + r"(?![a-z0-9+#])"
        if re.search(pattern, lower):
            skills.append(skill)
    education_items = extract_sections(text, "education")
    if not education_items:
        education_items = re.findall(r"(?i)(bachelor|master|phd|b\.s\.|m\.s\.|degree|university|college).{0,120}", compact)
    education = [{"description": item} for item in education_items[:5]]
    experience_items = extract_sections(text, "experience|work experience|employment")
    experience = [{"description": item, "years": extract_years(item)} for item in experience_items[:6]]
    project_items = extract_sections(text, "projects")
    projects = [{"title": item.split(".")[0][:80], "description": item, "technologies": [s for s in skills if s.lower() in item.lower()]} for item in project_items[:6]]
    cert_items = extract_sections(text, "certifications|certificates")
    certifications = cert_items[:8]
    return {
        "name": extract_name(lines, compact),
        "email": email_match.group(0) if email_match else "",
        "phone": phone_match.group(0) if phone_match else "",
        "skills": skills,
        "education": education,
        "experience": experience,
        "projects": projects,
        "certifications": certifications,
        "raw_text": raw_text,
        "word_count": len(compact.split()),
    }


def extract_years(text: str) -> float:
    total = 0.0
    for match in re.finditer(r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs)", text, flags=re.I):
        total = max(total, float(match.group(1)))
    years = [int(y) for y in re.findall(r"\b(20\d{2}|19\d{2})\b", text)]
    if len(years) >= 2:
        total = max(total, max(years) - min(years))
    return total


def parse_resume(file_or_path, filename: str | None = None) -> dict:
    if isinstance(file_or_path, (str, Path)):
        path = Path(file_or_path)
        suffix = path.suffix.lower()
        text = extract_pdf(path) if suffix == ".pdf" else extract_docx(path) if suffix == ".docx" else path.read_text(encoding="utf-8")
    else:
        if not filename:
            raise ValueError("filename is required for uploaded files")
        text = text_from_upload(file_or_path, filename)
    return parse_resume_text(text)
